import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_script(doc, text):
    p = doc.add_paragraph()
    p.style = 'Quote'
    run = p.add_run("🎙️ Script:\n" + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0, 51, 153) # Dark blue for script
    return p

doc = docx.Document()

# Title
title = doc.add_heading('KỊCH BẢN THUYẾT TRÌNH - X-AURUM', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Hệ Thống Dự Báo Giá Vàng Theo Thời Gian Thực', style='Subtitle')

p = doc.add_paragraph()
p.add_run('Members: ').bold = True
p.add_run('Xuân Hương · Quốc Đạt · Văn Nguyễn · Thành Trung')

doc.add_page_break()

# ----------------- XUÂN HƯƠNG -----------------
add_heading(doc, 'Phần 1: Xuân Hương (Từ Slide 1 - Slide 3)', 1)

add_heading(doc, 'Slide 1: Bài Toán & Kiến Trúc Cốt Lõi', 2)
add_script(doc, "Kính chào thầy cô và các bạn. Nhóm 3 chúng em xin bắt đầu bài báo cáo kết thúc môn với chủ đề: Triển khai Web API tích hợp Mô hình Trí tuệ Nhân tạo và dữ liệu giá Vàng thời gian thực. Hệ thống có tên X-AURUM. Em là Xuân Hương, em xin phép đi vào phần đầu tiên.\n\n"
                  "Bài toán đầu tiên nhóm em gặp phải là: Mô hình AI chỉ chạy trên máy cục bộ, không ai ngoài nhóm có thể dùng được. Để giải quyết, chúng em đã bọc toàn bộ logic này bên trong một Web API. Framework được chọn là ASP.NET Core Minimal API trên .NET 10 — đây là kiến trúc Microservices tinh gọn nhất hiện nay, loại bỏ hoàn toàn tầng Controller truyền thống phức tạp. Toàn bộ code được phân tách sạch sẽ thành 3 nhánh: Endpoints định tuyến, Models khai báo dữ liệu, và Extensions quản lý cấu hình.")

add_heading(doc, 'Slide 2: Thiết Kế Hai Luồng Endpoint', 2)
add_script(doc, "Nhóm em thiết kế 2 luồng endpoint độc lập. Luồng thứ nhất là POST — dành cho các hệ thống bên ngoài muốn tự truyền vào các chỉ số Open, High, Low, Volume để nhận dự báo.\n\n"
                  "Luồng thứ hai, và đây là điểm đặc biệt: GET /realtime — hoàn toàn không nhận bất kỳ input nào. Chính Server sẽ tự động đi lấy dữ liệu thị trường về và đưa vào mô hình. Mọi Response đều được chuẩn hóa chung một Wrapper JSON có cờ 'success', đảm bảo hệ thống gọi đến parse dữ liệu dễ dàng và an toàn.")

add_heading(doc, 'Slide 3: Tài Liệu API Tự Động Với Swagger (OpenAPI)', 2)
add_script(doc, "Hơn nữa, một API chuẩn doanh nghiệp không thể thiếu tài liệu mô tả. Thay vì viết tay file hướng dẫn, nhóm em tích hợp luôn Swagger — chuẩn mô tả OpenAPI tiên tiến nhất. Khi Server hoạt động, Swagger sẽ tự sinh trang web mô tả toàn bộ endpoints, kiểu dữ liệu vào-ra, và còn cho phép người dùng chạy thử (Execute) ngay trên web mà không cần cài thêm công cụ như Postman.\n\n"
                  "Lát nữa nhóm em cũng sẽ dùng Swagger để demo. Em xin kết thúc phần kiến trúc, tiếp theo mời bạn Quốc Đạt.")

# ----------------- QUỐC ĐẠT -----------------
add_heading(doc, 'Phần 2: Quốc Đạt (Từ Slide 4 - Slide 6)', 1)

add_heading(doc, 'Slide 4: Tích Hợp API Bên Thứ 3 (Binance & ExchangeRate)', 2)
add_script(doc, "Cảm ơn bạn Hương. Kính thưa thầy cô, em là Quốc Đạt, em xin trình bày về phần tích hợp 3rd-Party APIs. Quay lại bài toán: Server lấy dữ liệu giá vàng ở đâu nếu người dùng không nhập vào? Nhóm em để Server đóng vai Client đi gọi ra ngoài.\n\n"
                  "Nhóm tự động truy xuất từ 2 nguồn thực tế: Binance để lấy giá nến Vàng PAXG theo ngày, và Open Exchange Rates để lấy tỷ giá USD sang VND thời gian thực. Sau đó lấy kết quả AI nhân với tỷ giá để ra được số tiền VNĐ hiện tại.")

add_heading(doc, 'Slide 5: Xử Lý Bất Đồng Bộ & Phân Tích JSON', 2)
add_script(doc, "Để luồng gọi này hiệu quả nhất, nhóm em sử dụng kỹ thuật lập trình bất đồng bộ async/await. Khi có yêu cầu, Server đồng thời khởi tạo 2 kết nối ra Binance và Exchange Rates. Khi đó luồng xử lý không hề bị chặn lại chờ đợi mạng, giúp Server tiết kiệm tài nguyên.\n\n"
                  "Sau khi nhận đủ dữ liệu, hệ thống dùng System.Text.Json đọc trực tiếp các Node cần thiết trên chuỗi JSON gửi về mà không cần ánh xạ (map) ra toàn bộ Object lớn, rất nhẹ và tiết kiệm RAM.")

add_heading(doc, 'Slide 6: IHttpClientFactory - Gọi API Chuyên Nghiệp', 2)
add_script(doc, "Một điểm mạnh nữa về kỹ thuật: nhóm em không khởi tạo 'new HttpClient' mỗi khi gọi ra ngoài, làm như vậy sẽ dẫn đến Socket Exhaustion — hiểu nôm na là cạn kiệt cổng mạng khi scale lớn.\n\n"
                  "Chúng em dùng IHttpClientFactory, cung cấp sẵn bởi .NET DI Container. Factory này tái sử dụng các kết nối HTTP vô cùng thông minh, không lo rò rỉ Socket và giúp ứng dụng chịu tải cao cực tốt. Phần tiếp theo về hiệu năng, xin mời bạn Văn Nguyễn.")

# ----------------- VĂN NGUYỄN -----------------
add_heading(doc, 'Phần 3: Văn Nguyễn (Từ Slide 7 - Slide 9)', 1)

add_heading(doc, 'Slide 7: Bài Toán Hiệu Năng - Hosting ML Model', 2)
add_script(doc, "Kính chào thầy cô, em là Văn Nguyễn, em phụ trách tối ưu hiệu năng và bảo mật hệ thống.\n\n"
                  "File AI học máy của nhóm nặng hơn 40MB. Thách thức lớn là: nếu mỗi lúc nhận Request, API lại đọc file 40MB này từ ổ cứng vào RAM thì sẽ mất hàng chục giây cho mỗi lượt tương tác, vài nghìn Request vào cùng lúc chắc chắn Server sẽ sập vì tràn RAM. Đoạn code sai lầm ở trên File PPT là thứ chúng em đã loại bỏ hoàn toàn.")

add_heading(doc, 'Slide 8: PredictionEnginePool - Quản Trị RAM', 2)
add_script(doc, "Cách chúng em giải quyết là nhúng PredictionEnginePool — một dạng Object Pooling của thư viện Microsoft.ML.\n\n"
                  "Thay vì load model liên tục, Pooling chỉ load model lên RAM đúng 1 lần duy nhất khi Server khởi động. Khi có Request, API chỉ mượn một Engine đang rảnh trong hồ bơi (Pool), dùng xong lại trả về. Cơ chế này an toàn về Thread (Thread-Safe) cho dù có hàng nghìn lượt Request đồng thời và loại bỏ hoàn toàn hiện tượng Memory Leak.")

add_heading(doc, 'Slide 9: Rate Limiting - Tường Lửa Chống Tấn Công', 2)
add_script(doc, "Bên cạnh tối ưu RAM, API của chúng em là công khai nên rất dễ bị phá hoại hoặc spam bằng tấn công DDoS. Nhóm em đã bật tường lửa Rate Limiting theo thuật toán Fixed Window. Mỗi IP sẽ chỉ được gọi 5 Request trong 10 giây.\n\n"
                  "Hãy tưởng tượng có kẻ xấu spam 1 triệu Request, nó sẽ bị chặn ngay tại tầng mạng HTTP Pipeline và trả về HTTP 429 Too Many Requests, hoàn toàn không đẩy lượng rác đó vào phân tích AI. Server được bảo vệ tuyệt đối. Tiếp theo mời bạn Thành Trung.")

# ----------------- THÀNH TRUNG -----------------
add_heading(doc, 'Phần 4: Thành Trung (Từ Slide 10 - Slide 12)', 1)

add_heading(doc, 'Slide 10: Docker - Đóng Gói & Chuẩn Hóa Môi Trường', 2)
add_script(doc, "Cảm ơn bạn Nguyễn. Thưa thầy cô, em là Thành Trung, em xin khép lại bài thuyết trình với kiến trúc triển khai DevOps thiết thực.\n\n"
                  "Lúc làm việc nhóm trong suốt kì học qua, một vấn đề cực kỳ khó chịu là 'Code máy em chạy nhưng copy sang máy bạn cài .NET bản khác là lỗi'. Chúng em chấm dứt điều này với Docker. \n"
                  "Chúng em đóng thùng toàn bộ Server vào Container, viết Dockerfile theo Multi-Stage Build — giai đoạn một dùng SDK cả Gb để Build, giai đoạn hai chuyển thành Runtime nhẹ tựa lông hồng, chỉ 200MB. Mọi hệ điều hành chạy đều đồng nhất bản chất, 1 lệnh duy nhất là hệ thống vận hành.")

add_heading(doc, 'Slide 11: CI/CD & Triển Khai Trên nttspace.online', 2)
add_script(doc, "Để tiết kiệm thời gian deploy thủ công, nhóm tích hợp Workflow GitHub Actions tự động hóa. Mỗi khi ấn Push Code vòng tuần hoàn sẽ diễn ra: Git Server tự Build Docker Image, đẩy về Registry, và Virtual Machine trên Cloud sẽ tự biết kéo Image mới về khởi động lại không gián đoạn dịch vụ.\n\n"
                  "Ngay hiện tại, hệ thống đã được Public Production tại domain nttspace.online với chứng chỉ bảo mật HTTPS đầy đủ, sử dụng Nginx Reverse Proxy làm mặt tiền hứng toàn bộ traffic để truyền vào bên trong an toàn.")

add_heading(doc, 'Slide 12: Live Demo Trên Swagger', 2)
add_script(doc, "Và sau đây, không có gì thuyết phục hơn là chạy demo thực tế. Trên bảng lớn là màn hình Swagger thật từ tên miền chính thức của nhóm là nttspace.online.\n\n"
                  "Em sẽ thực thi lệnh GET Realtime không cấp dữ liệu — màn hình tức thời hiển thị giá vàng thế giới và đổi ra VND trực tiếp cực kì chính xác.\n\n"
                  "Kế đến, em xin spam nút Execute liên tục vào Endpoint POST. Như bạn Nguyễn đã thiết kế phần trước, thầy cô có thể thấy màn hình bật qua trạng thái màu đỏ 429 Too Many Requests — tường lửa đã thành công khóa địa chỉ IP của chính em lại. \n\n"
                  "Đó là tất cả tính năng của dự án X-AURUM phục vụ cho bài thi lần này. Thay mặt Nhóm 3 em xin cảm ơn thầy cô đã lắng nghe và chúng em xin giải đáp các câu hỏi phản biện ạ!")

doc.save('Kich_Ban_Thuyet_Trinh_X-AURUM.docx')
print("Saved DOCX file!")
